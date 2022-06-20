import json
import string

import docx
import pymongo as pymongo
import redis as redis
from docx import Document
from flask import Flask, render_template, jsonify, redirect, url_for, request
from redis import StrictRedis
from redis.commands.json.path import Path

app = Flask(__name__)


try:
    redis = redis.Redis(
        host='localhost',
        port='6379',
        charset='utf-8',
        decode_responses=True
    )

    mongo = pymongo.MongoClient(
        host ="localhost",
        port=27017,
        serverSelectionTimeoutMS = 1000
    )
    db = mongo.QCMQUIZ
    mongo.server_info()
    print("Connexion établie")
except:
    print("ERREUR-Impossibe de se connecter à la BDD")


@app.route('/')
def welcome_page():
    return render_template('index.html', name='dashbord')

@app.route('/addQuestions')
def add_question():
    return render_template('addQuestions.html', name='ajouter des questions')

@app.route('/allQuestions')
def all_question():
    #mongo.ge
    return render_template('allQuestions.html', name='afficher les questions')


@app.route('/updateQuestion', methods=['POST'])
def update():
    print("toto")
    print(list(request.form))
    return render_template("index.html")

@app.route('/deleteQuestion/<id>')
def toto(id):
    redis.delete(f"questions:{id}")

    questionskeys = redis.keys("*")

    listKeys = list(questionskeys)
    listKeys = sorted(listKeys)

    print(listKeys)
    listnum = []
    for item in listKeys:
        listnum.append(int(item[10:]))
    listnum = sorted(listnum)
    questions = []

    for i in listnum:
        cle = "questions:" + str(i)
        questions.append(redis.hgetall(cle))
        #print("je suis dans la boucle")
        #print(i)
    return render_template('questionsRedis.html', data=questions,  len = len(questions))


@app.route('/deleteQuestion/<id>')
def delete_question(id):
    print(f"***** je passe par la question {id}")
    redis.delete(f"questions:{id}")
    nbQuestions = nbQuestionsRedis()
    questions = []
    for i in range(nbQuestions):
        questions.append(redis.hgetall("questions:" + i.__str__()))
    #while {} in questions:
        #questions.remove({})

    return render_template('questionsRedis.html', data=questions,  len = len(questions))


@app.route('/questionsJSON')
def getALLQuestionsRedis():
    print("*********")
    print (str(nbQuestionsRedis()))
    questions = []
    for i in range(nbQuestionsRedis()):
        print (str(i))
        questions.append(redis.hgetall("questions:" + i.__str__()))
        print(redis.hgetall("questions:" + i.__str__()))
        app.config['JSON_AS_ASCII'] = False
    while {} in questions:
        questions.remove({})
    return questions

def nbQuestionsRedis():
    print("nombre de questions: ")
    print(len(redis.keys("*")))
    return len(redis.keys("*"))

@app.route('/q')
def questions_page():
    nbQuestions = nbQuestionsRedis()
    import_questions_redis()
    questions =[]
    for i in range(nbQuestions):
        questions.append(redis.hgetall("questions:" + i.__str__()))
    return json.dumps(questions)


@app.route('/questions')
def json_questions():

    questions = getALLQuestionsRedis()
    nbQuestions = nbQuestionsRedis()
    return render_template('questionsRedis.html', data=questions,  len = nbQuestions)

@app.route('/questionsRedis')
def import_questions_redis():
    correction = ocr_World_Correction()
    questions = ocr_World_Questions()
    correction.pop(0)
    for i in range(len(questions)):
        questions[i].append(correction[i])
        redisData = {}
        redisData["question"] = questions[i][0]
        redisData["optionA"] = questions[i][1]
        redisData["optionB"] = questions[i][2]
        redisData["optionC"] = questions[i][3]
        redisData["optionD"] = questions[i][4]
        redisData["numero"] = i
        if (questions[i][5][1] == "x"):
            correctAnswer = 1
        elif(questions[i][5][2] == "x"):
            correctAnswer = 2
        elif(questions[i][5][3] == "x"):
            correctAnswer = 3
        elif (questions[i][5][4] == "x"):
            correctAnswer = 4
        redisData["correctOption"] = correctAnswer
        redis.hmset("questions:"+i.__str__(), redisData)

    return i


@app.route('/ocrWordCorrection')
def ocr_World_Correction():
    document = Document('test.docx')
    responses = []
    for row in document.tables[0].rows:
        el = []
        for cell in row.cells:
            for para in cell.paragraphs:
                el.append((para.text).strip())
        responses.append(el)


    return responses

@app.route('/ocrWordQuestions')
def ocr_World_Questions():
    doc = docx.Document('sujet.docx')  # Creating word reader object.
    data = ""
    fullText = []
    nbPoints = 0;
    for para in doc.paragraphs:
        fullText.append(para.text)
        data = '\n'.join(fullText)
    debut = data.index("1/ QCM (1h30)")
    fin = data.index("2/ Epreuve écrite (2h30)")
    data = data[debut+15:fin]
    data = data.strip()
    data = data.split("\n")
    data.remove(" ")
    data.remove("")
    data.remove(' ')
    data.remove('')
    while '' in data:
        del data[data.index('')]
    while " " in data:
        del data[data.index(" ")]

    n=5

    output = [data[i:i + n] for i in range(0, len(data), n)]


    return  output

def upload_questions():
    questions = getALLQuestionsRedis()

    return questions
if __name__ == '__main__':
    app.run()
    ocr_World_Questions()
enctype="multipart/form-data"
# importing the required libraries
from flask import Flask, render_template, request
from werkzeug.utils import secure_filename

# initialising the flask app
app = Flask(__name__)


# The path for uploading the file
@app.route('/')
def upload_file():
    return render_template('upload.html')


@app.route('/upload', methods=['GET', 'POST'])
def uploadfile():
    if request.method == 'POST':  # check if the method is post
        f = request.files['file']  # get the file from the files object
        f.save(secure_filename(f.filename))  # this will secure the file
        return 'file uploaded successfully'  # Display thsi message after uploading


if __name__ == '__main__':
    app.run()  # running the flask app