import json
import os
import random
import string
import time

import flask
from werkzeug.utils import secure_filename

import correction

import bson
import docx
import pymongo as pymongo
import redis as redis
from bson import json_util
from docx import Document
from docx.oxml import OxmlElement, ns
from docx.shared import Inches
from flask import Flask, render_template, jsonify, redirect, url_for, request, flash
from redis import StrictRedis
from redis.commands.json.path import Path

app = Flask(__name__)

app.secret_key = "secret key"

try:
    #connection to redis database
    redis = redis.Redis(
        host='localhost',
        port='6379',
        charset='utf-8',
        decode_responses=True
    )

    #connection to mongodb database
    mongo = pymongo.MongoClient(
        host ="localhost",
        port=27017,
        serverSelectionTimeoutMS = 1000

    )
    db = mongo["QCMQUIZ"]
    #print(mongo.server_info())


    print("Connexion établie")
except:
    print("ERREUR-Impossibe de se connecter à la BDD")


@app.route('/')
def welcome_page():
    return render_template('index.html', name='dashbord')

#affichage de la page de génération d'un nouveau sujet
@app.route('/nouveau-sujet')
def genQCTemplate():
    return render_template('genQCM.html', name='genQCM')

#génération d'un nouveau sujet
@app.route('/newQCM', methods=['POST'])
def genQCM():
    listQuestions = {"matiere": request.form["matiere"], "niveau": request.form["niveau"]}
    # Collection name
    collection = db[request.form["matiere"]]

    # if we don't want to print id then pass _id:0
    #for x in collection.find({}, {"_id": 0, "question": 1, "correctOption": 1, "optionA": 1, "optionB": 1, "optionC": 1, "optionD":1 }).limit(int(request.form["nbQuestions"])):

    x = collection.aggregate([{"$project": {"_id": 0, "question": 1, "correctOption": 1, "optionA": 1, "optionB": 1,
                                        "optionC": 1, "optionD": 1, "numero": 1}},
                          {"$sample": {"size": int(request.form["nbQuestions"])}}])

    listQuestions["questions"] = list(x)

    #print(listQuestions)

    def create_element(name):
        return OxmlElement(name)

    def create_attribute(element, name, value):
        element.set(ns.qn(name), value)

    def add_page_number(run):
        fldChar1 = create_element('w:fldChar')
        create_attribute(fldChar1, 'w:fldCharType', 'begin')

        instrText = create_element('w:instrText')
        create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"

        fldChar2 = create_element('w:fldChar')
        create_attribute(fldChar2, 'w:fldCharType', 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
    document = Document()
    document2 = Document()

    document.add_picture('unnamed.jpg', width=Inches(1.25))

    document.add_heading('Sujet: '+request.form["matiere"]+' '+request.form["niveau"], 0)
    document2.add_heading(
        'Correction du sujet: ' + request.form["matiere"] + ' ' + request.form["niveau"] + ' du ' + time.strftime(
            '%d-%m-%Y'))
    document.add_heading('Consignes: ', 2)
    document.add_paragraph('Pour chaque question, une seule réponse est bonne.')
    document.add_paragraph('Une bonne réponse vaut 1 point, une mauvaise vaut 0.')
    document.add_paragraph('Dans la fiche des réponses, cochez les cases en mettant un « X » pour que la réponse soit prise en compte.')
    document.add_heading('Questions: ', 2)
    #création du tableau des réponses
    table = document2.add_table(rows=len(listQuestions)-2, cols=5, style='TableGrid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = ''
    hdr_cells[1].text = 'A'
    hdr_cells[2].text = 'B'
    hdr_cells[3].text = 'C'
    hdr_cells[4].text = 'D'
    i=1


    # Adding points to the list named 'List Number'
    numero = 0
    for item in listQuestions["questions"]:
        print("*********")
        print(item)
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        if(item['correctOption'] == "1") : row_cells[1].text = "x"
        if(item['correctOption'] == "2") : row_cells[2].text = "x"
        if(item['correctOption'] == "3") : row_cells[3].text = "x"
        if(item['correctOption'] == "4") : row_cells[4].text = "x"
        i += 1
        numero +=1
        #print(item["question"])
        # Adding list of style name 'List Bullet'
        document.add_heading('Question '+str(numero)+ ': ' + item['question'], 3)

        document.add_paragraph('A: ' + item['optionA'],
                               style='List Bullet 2')
        document.add_paragraph('B: ' + item['optionB'],
                               style='List Bullet 2')
        document.add_paragraph('C: ' + item['optionC'],
                               style='List Bullet 2')
        document.add_paragraph('D: ' + item['optionD'],
                               style='List Bullet 2')
    add_page_number(document.sections[0].footer.paragraphs[0].add_run())

    document.save('sujet_'+request.form["matiere"].replace(".", "-")+'_'+request.form["niveau"]+'_'+time.strftime('%d-%m-%Y')+'.docx')
    document2.save('Correction_sujet_'+request.form["matiere"].replace(".", "-")+'_'+request.form["niveau"]+'_'+time.strftime('%d-%m-%Y')+'.docx')

    return  json.dumps(listQuestions, ensure_ascii=False)


#Alimentation de la base mongodb
@app.route('/upload_questions_mongo')
def upload_questions_mongo():
    questions = getALLQuestionsRedis()
    collection = db['D51.1']

    for question in questions:
        collection.insert_one(question)
    return render_template('addQuestions.html', name='ajouter des questions')

#Affichage de toutes les questions d'une matière passée en paramètre
@app.route('/getallQuestions')
def getALLQuestions():
    #collection = db['D51.1']
    collection = db['D51.1']
    documents = list(collection.find())


    return render_template('allQuestions.html', name='afficher les questions')

@app.route('/allQuestions')
def all_question():
    questions = getALLQuestions()

    return render_template('allQuestions.html', name='afficher les questions')

#modification d'une question en REDIS
@app.route('/updateQuestion', methods=['POST', 'GET'])
def update():

    dicQuestion = {'question': '', 'optionA': '', 'optionB': '', 'optionC': '', 'optionD': '', 'numero': '', 'correctOption': ''}
    dicQuestion["question"] = request.form["question"]
    dicQuestion["optionA"] = request.form["optionA"]
    dicQuestion["optionB"] = request.form["optionB"]
    dicQuestion["optionC"] = request.form["optionC"]
    dicQuestion["optionD"] = request.form["optionD"]
    dicQuestion["correctOption"] = request.form["correctAnswer"]
    dicQuestion["numero"] = request.form["id"]
    redis.delete(f"questions:{request.form['id']}")
    redis.hmset(f"questions:{request.form['id']}", dicQuestion)
    return redirect("/questions")

#Suppression d'une question en REDIS
@app.route('/deleteQuestion/<id>')
def toto(id):
    redis.delete(f"questions:{id}")
    return redirect('/questions')

#Récupérations de tous les questions en REDIS
def getALLQuestionsRedis():
    questions = []
    for i in range(nbQuestionsRedis()):
        questions.append(redis.hgetall("questions:" + i.__str__()))
        app.config['JSON_AS_ASCII'] = False
    while {} in questions:
        questions.remove({})
    return questions

#Retournr le nombre de questions en REDIS
def nbQuestionsRedis():
    return len(redis.keys("*"))

@app.route('/q')
def questions_page():
    nbQuestions = nbQuestionsRedis()
    import_questions_redis()
    questions =[]
    for i in range(nbQuestions):
        questions.append(redis.hgetall("questions:" + i.__str__()))
    return json.dumps(questions)

#
@app.route('/parser-un-sujet')
def parser_document():
    return render_template("newParse.html")

@app.route('/corriger-examen')
def corriger_exam():
    return render_template("corriger.html")\
# Get current path
path = os.getcwd()
# file Upload
UPLOAD_FOLDER = os.path.join(path, 'uploads')

# Make directory if uploads is not exists
if not os.path.isdir(UPLOAD_FOLDER):
    os.mkdir(UPLOAD_FOLDER)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

@app.route('/correctionExam', methods=["GET","POST"])
def correctExamModel():
    if request.method == 'POST':

        if 'answerFile[]' not in request.files:
            flash('No file part')
            return redirect(request.url)

        files = request.files.getlist('answerFile[]')

        for file in files:
            if file :
                filename = secure_filename(file.filename)
                file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))

        flash('File(s) successfully uploaded')
        return redirect('/')



   # correction.corriger_qcm(correctionFile, answerFiles)
    return json.dumps("file uploaded")


@app.route('/questionsParsed', methods=['POST', 'GET'])
def json_questions2():
    import_questions_redis()
    #import_questions_redis2(request.form['fileCorrection'], request.form['fileSujet'])

    questionskeys = redis.keys("*")

    listKeys = list(questionskeys)
    listKeys = sorted(listKeys)

    listnum = []
    for item in listKeys:
        listnum.append(int(item[10:]))
    listnum = sorted(listnum)
    questions = []

    for i in listnum:
        cle = "questions:" + str(i)
        questions.append(redis.hgetall(cle))

    return render_template('questionsRedis.html', data=questions,  len = len(questions))


@app.route('/questions', methods=['POST', 'GET'])

def ff():
    questions = getALLQuestionsRedis()
    return render_template('questionsRedis.html', data=questions,  len = len(questions))

@app.route('/questionsRedis')
def import_questions_redis():
    correction = ocr_World_Correction()
    questions = ocr_World_Questions()
    correction.pop(0)
    for i in range(len(questions)):
        questions[i].append(correction[i])
        redisData = {}
        redisData["question"] = questions[i][0]
        redisData["optionA"] = questions[i][1]
        redisData["optionB"] = questions[i][2]
        redisData["optionC"] = questions[i][3]
        redisData["optionD"] = questions[i][4]
        redisData["numero"] = i
        if (questions[i][5][1] == "x"):
            correctAnswer = 1
        elif(questions[i][5][2] == "x"):
            correctAnswer = 2
        elif(questions[i][5][3] == "x"):
            correctAnswer = 3
        elif (questions[i][5][4] == "x"):
            correctAnswer = 4
        redisData["correctOption"] = correctAnswer
        redis.hmset("questions:"+i.__str__(), redisData)

    return i



@app.route('/questionsRedis2')
def import_questions_redis2(fileCorrection, fileSujet):
    correction = ocr_World_Correction2("test.docx")
    questions = ocr_World_Questions2("sujet2.docx")
    correction.pop(0)
    for i in range(len(questions)):
        questions[i].append(correction[i])
        redisData = {}
        redisData["question"] = questions[i][0]
        redisData["optionA"] = questions[i][1]
        redisData["optionB"] = questions[i][2]
        redisData["optionC"] = questions[i][3]
        redisData["optionD"] = questions[i][4]
        redisData["numero"] = i
        if (questions[i][5][1] == "x"):
            correctAnswer = 1
        elif(questions[i][5][2] == "x"):
            correctAnswer = 2
        elif(questions[i][5][3] == "x"):
            correctAnswer = 3
        elif (questions[i][5][4] == "x"):
            correctAnswer = 4
        redisData["correctOption"] = correctAnswer
        redis.hmset("questions:"+i.__str__(), redisData)
    return i
@app.route('/ocrWordCorrection')
def ocr_World_Correction():
    document = Document('test.docx')
    responses = []
    for row in document.tables[0].rows:
        el = []
        for cell in row.cells:
            for para in cell.paragraphs:
                el.append((para.text).strip())
        responses.append(el)

    return responses

@app.route('/ocrWordQuestions')
def ocr_World_Questions():
    doc = docx.Document('sujet2.docx')  # Creating word reader object.
    data = ""
    fullText = []
    nbPoints = 0;
    for para in doc.paragraphs:
        fullText.append(para.text)
        data = '\n'.join(fullText)
    debut = data.index("1/ QCM (1h30)")
    fin = data.index("2/ Epreuve écrite (2h30)")
    data = data[debut+15:fin]
    data = data.strip()
    data = data.split("\n")
    data.remove(" ")
    data.remove("")
    data.remove(' ')
    data.remove('')
    while '' in data:
        del data[data.index('')]
    while " " in data:
        del data[data.index(" ")]

    n=5

    output = [data[i:i + n] for i in range(0, len(data), n)]


    return  output


@app.route('/ocrWordCorrection2')
def ocr_World_Correction2(fileCorrection):
    document = Document("test.docx")
    responses = []
    for row in document.tables[0].rows:
        el = []
        for cell in row.cells:
            for para in cell.paragraphs:
                el.append((para.text).strip())
        responses.append(el)

    return responses

@app.route('/ocrWordQuestions2')
def ocr_World_Questions2(fileSujet):
    doc = docx.Document("sujet2.docx")  # Creating word reader object.
    data = ""
    fullText = []
    nbPoints = 0;
    for para in doc.paragraphs:
        fullText.append(para.text)
        data = '\n'.join(fullText)
    debut = data.index("1/ QCM (1h30)")
    fin = data.index("2/ Epreuve écrite (2h30)")
    data = data[debut+15:fin]
    data = data.strip()
    data = data.split("\n")
    data.remove(" ")
    data.remove("")
    data.remove(' ')
    data.remove('')
    while '' in data:
        del data[data.index('')]
    while " " in data:
        del data[data.index(" ")]

    n=5

    output = [data[i:i + n] for i in range(0, len(data), n)]


    return  output

def upload_questions():
    questions = getALLQuestionsRedis()

    return questions
if __name__ == '__main__':
    correction.listFiles('h:/')
    app.run()
